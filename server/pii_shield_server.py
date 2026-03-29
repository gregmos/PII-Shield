"""
PII Shield MCP Server v5.4.0
==============================
Cowork-only. NER-only mode (no LLM dependency).
  - NER backend: Presidio TransformersNlpEngine (dslim/bert-base-NER) with SpaCy tokenization
  - High-quality BERT NER via Presidio's native transformer pipeline
  - Self-bootstrapping: auto-installs missing packages on first run
  - DRY: single detect() used by both text and docx paths
  - TTL: mappings auto-cleaned after 7 days
  - Unified anonymize_file: routes .docx/.txt/.pdf automatically
  - EU recognizers: 17 patterns (UK, DE, FR, IT, ES, CY, EU-wide)
  - Indexed placeholders: <PERSON_1>, <ORG_1> with reversible mapping
  - Fuzzy entity deduplication, prefix support, boundary cleanup
  - PII-safe responses: mapping and real text never returned to Claude

Tools:
  anonymize_text / anonymize_file / anonymize_docx
  deanonymize_text / deanonymize_docx
  get_mapping / scan_text / list_entities
"""

# ============================================================
# Self-bootstrap: three-phase progressive installation
#
#   Phase 1 (synchronous, ~2s):  install ONLY mcp — so the server can start
#   Phase 2 (background thread): install all heavy packages (torch, presidio, etc.)
#   Phase 3 (background thread): download AI models (spacy, bert-base-NER)
#
# Server starts accepting MCP connections after Phase 1 (~2 seconds).
# Tools respond with installation progress until Phase 2+3 complete.
# No timeouts, no manual setup scripts needed.
# ============================================================
import subprocess
import sys
import threading
import logging as _boot_log

_boot_log.basicConfig(level=_boot_log.INFO, format="%(asctime)s [PII-Shield] %(message)s", stream=sys.stderr)
_blog = _boot_log.getLogger("pii-shield-bootstrap")

_GLINER_MODEL = "urchade/gliner_small-v2.1"

# Bootstrap state — read by tools to report progress
_bootstrap_phase = "starting"    # starting → packages → models → ready / error
_bootstrap_detail = ""           # human-readable progress detail
_bootstrap_done = False
_bootstrap_error = None
_bootstrap_start_time = None     # set when bootstrap begins

# --- Status file for Cowork Skill warm-up ---
import json as _json_boot
from pathlib import Path as _BootPath

_STATUS_DIR = _BootPath.home() / ".pii_shield"
_STATUS_FILE = _STATUS_DIR / "status.json"


def _write_status(phase, detail="", progress_pct=0):
    """Write bootstrap status to ~/.pii_shield/status.json for Skill warm-up monitoring."""
    import time as _t
    try:
        _STATUS_DIR.mkdir(parents=True, exist_ok=True)
        elapsed = round(_t.time() - _bootstrap_start_time, 1) if _bootstrap_start_time else 0
        _STATUS_FILE.write_text(_json_boot.dumps({
            "phase": phase,
            "message": detail,
            "progress_pct": progress_pct,
            "elapsed_seconds": elapsed,
            "timestamp": _t.time(),
        }, indent=2), encoding="utf-8")
    except Exception:
        pass  # non-critical, don't crash bootstrap

# Packages split into phases: MCP first (tiny), then everything else (heavy)
_MCP_PACKAGE = ("mcp", "mcp[cli]>=1.0.0")
_HEAVY_PACKAGES = [
    ("presidio_analyzer", "presidio-analyzer>=2.2.355"),
    ("spacy",             "spacy>=3.7.0"),
    ("docx",              "python-docx>=1.1.0"),
    ("cryptography",      "cryptography>=42.0.0"),
    ("numpy",             "numpy>=1.24.0"),
    ("torch",             "torch>=2.0.0"),
    ("gliner",            "gliner>=0.2.7"),
]


def _install_if_missing(packages):
    """Install missing pip packages. Returns list of installed specs."""
    missing = []
    for import_name, pip_spec in packages:
        try:
            __import__(import_name)
        except ImportError:
            missing.append(pip_spec)
    if missing:
        _blog.info(f"Installing: {missing}")
        # Fully detach pip from parent's stdio to prevent hanging in Cowork.
        # On Windows, CREATE_NO_WINDOW prevents console inheritance issues.
        import os as _os
        _pip_dir = _os.path.join(_os.path.expanduser("~"), ".pii_shield")
        _os.makedirs(_pip_dir, exist_ok=True)
        _pip_log = _os.path.join(_pip_dir, "pip_install.log")
        _cflags = 0
        if sys.platform == "win32":
            _cflags = subprocess.CREATE_NO_WINDOW
        with open(_pip_log, "a") as _lf:
            subprocess.check_call(
                [sys.executable, "-m", "pip", "install", "--quiet"] + missing,
                stdin=subprocess.DEVNULL, stdout=_lf, stderr=_lf,
                creationflags=_cflags,
            )
    return missing


def _download_models():
    """Download SpaCy + transformer models if not cached."""
    import os

    # SpaCy model
    try:
        import spacy
        try:
            spacy.load("en_core_web_sm")
            _blog.info("SpaCy model: cached")
        except OSError:
            _blog.info("SpaCy model: downloading...")
            _pip_log = os.path.join(os.path.expanduser("~"), ".pii_shield", "pip_install.log")
            with open(_pip_log, "a") as _lf:
                _cflags = 0
                if sys.platform == "win32":
                    _cflags = subprocess.CREATE_NO_WINDOW
                subprocess.check_call(
                    [sys.executable, "-m", "spacy", "download", "en_core_web_sm"],
                    stdin=subprocess.DEVNULL, stdout=_lf, stderr=_lf,
                    creationflags=_cflags,
                )
    except Exception as e:
        _blog.warning(f"SpaCy model: {e}")

    # GLiNER model (with retries)
    try:
        gliner_model = os.environ.get("PII_GLINER_MODEL", _GLINER_MODEL)
        _blog.info(f"GLiNER model {gliner_model}: checking cache...")
        _write_status("models", f"Downloading GLiNER NER model ({gliner_model})... ~900 MB", 75)
        from gliner import GLiNER
        _retries = [0, 10, 30, 60]
        for _attempt, _delay in enumerate(_retries):
            if _delay > 0:
                _blog.info(f"Retry {_attempt}/{len(_retries)-1} in {_delay}s...")
                import time as _t; _t.sleep(_delay)
            try:
                GLiNER.from_pretrained(gliner_model)
                _blog.info(f"GLiNER model {gliner_model}: ready")
                break
            except Exception as _dl_err:
                _blog.warning(f"GLiNER download attempt {_attempt+1} failed: {_dl_err}")
                if _attempt == len(_retries) - 1:
                    raise
    except Exception as e:
        _blog.warning(f"GLiNER model download failed (will retry on first use): {e}")


# --- Phase 1 (synchronous): install ONLY mcp so FastMCP server can start ---
# With CREATE_NO_WINDOW + stdin=DEVNULL, pip won't hang in Cowork.
# Takes ~20s on first install; Cowork waits up to 60s for handshake.
import time as _time_boot
_bootstrap_start_time = _time_boot.time()
_write_status("starting", "PII Shield is starting up...")
try:
    _install_if_missing([_MCP_PACKAGE])
except Exception as _e:
    _bootstrap_error = str(_e)
    _blog.error(f"MCP install failed: {_e}")
    _write_status("error", f"MCP install failed: {_e}")


def _background_bootstrap():
    """Phase 2+3: install heavy packages and download models in background."""
    global _bootstrap_phase, _bootstrap_detail, _bootstrap_done, _bootstrap_error
    try:
        # Phase 2: check/install heavy packages
        _bootstrap_phase = "packages"
        _blog.info("Phase 2: checking packages...")
        _write_status("packages", "Checking dependencies...", 10)
        installed = _install_if_missing(_HEAVY_PACKAGES)
        if installed:
            _bootstrap_detail = "Installing dependencies (PyTorch, Presidio, SpaCy)... This takes 5-10 min on first run."
            _write_status("packages", _bootstrap_detail, 20)
            _blog.info(f"Phase 2 complete: installed {len(installed)} packages")
        else:
            _blog.info("Phase 2: all packages already installed")

        # Phase 3: load models into memory (required — GLiNER takes ~30-60s to load)
        _bootstrap_phase = "models"
        _bootstrap_detail = "Loading AI models into memory (~30-60s)..."
        _write_status("models", _bootstrap_detail, 70)
        _blog.info("Phase 3: loading models...")
        _download_models()
        _blog.info("Phase 3 complete")

        # Phase 4: initialize the PII engine (Presidio + GLiNER recognizer)
        # This avoids 60s+ timeout on first tool call.
        _bootstrap_phase = "engine"
        _bootstrap_detail = "Initializing PII engine..."
        _write_status("engine", _bootstrap_detail, 90)
        _blog.info("Phase 4: initializing PII engine...")
        # Wait for module-level `engine = PIIEngine()` to be defined
        for _w in range(30):
            if "engine" in globals():
                break
            import time as _tw; _tw.sleep(1)
        if "engine" in globals():
            engine._ensure_ready(_from_bootstrap=True)
            _blog.info("Phase 4 complete — engine initialized")
        else:
            _blog.warning("Phase 4: engine not yet defined, will init on first tool call")

        _bootstrap_phase = "ready"
        _bootstrap_detail = ""
        _write_status("ready", "PII Shield is ready.", 100)
        _blog.info("Bootstrap complete — ready for tool calls")

    except Exception as e:
        _bootstrap_error = str(e)
        _bootstrap_phase = "error"
        _bootstrap_detail = f"Bootstrap failed: {e}"
        _write_status("error", _bootstrap_detail, 0)
        _blog.error(f"Bootstrap failed: {e}")
    finally:
        _bootstrap_done = True

# --- Phase 2+3 (background): everything else ---
_bg_thread = threading.Thread(target=_background_bootstrap, daemon=True)
_bg_thread.start()
# ============================================================

import json
import os
import re
import time
import uuid
import logging
from pathlib import Path
from collections import defaultdict

from mcp.server.fastmcp import FastMCP
# NOTE: presidio imports are deferred to _ensure_ready() — not yet installed on first run

logging.basicConfig(level=logging.INFO, format="%(asctime)s [PII-Shield] %(message)s", stream=sys.stderr)
log = logging.getLogger("pii-shield")

# ============================================================
# Config
# ============================================================
MIN_SCORE = float(os.environ.get("PII_MIN_SCORE", "0.35"))
MAPPING_TTL_DAYS = int(os.environ.get("PII_MAPPING_TTL_DAYS", "7"))

MAPPING_DIR = Path.home() / ".pii_shield" / "mappings"
try:
    MAPPING_DIR.mkdir(parents=True, exist_ok=True)
except Exception:
    pass  # will retry in save_mapping; in-memory fallback always works

SUPPORTED_ENTITIES = [
    "PERSON", "ORGANIZATION", "LOCATION", "DATE_TIME", "NRP",
    "EMAIL_ADDRESS", "PHONE_NUMBER", "URL", "IP_ADDRESS",
    "CREDIT_CARD", "IBAN_CODE", "CRYPTO",
    "US_SSN", "US_PASSPORT", "US_DRIVER_LICENSE",
    "UK_NHS", "UK_NIN", "UK_PASSPORT", "UK_CRN", "UK_DRIVING_LICENCE",
    "EU_VAT", "EU_PASSPORT",
    "DE_TAX_ID", "DE_SOCIAL_SECURITY",
    "FR_NIR", "FR_CNI",
    "IT_FISCAL_CODE", "IT_VAT",
    "ES_DNI", "ES_NIE",
    "CY_TIC", "CY_ID_CARD",
    "MEDICAL_LICENSE",
]

TAG_NAMES = {
    "PERSON": "PERSON", "ORGANIZATION": "ORG", "LOCATION": "LOCATION",
    "DATE_TIME": "DATE", "NRP": "NRP",
    "EMAIL_ADDRESS": "EMAIL", "PHONE_NUMBER": "PHONE", "URL": "URL",
    "IP_ADDRESS": "IP", "CREDIT_CARD": "CREDIT_CARD", "IBAN_CODE": "IBAN",
    "CRYPTO": "CRYPTO",
    "US_SSN": "US_SSN", "US_PASSPORT": "US_PASSPORT", "US_DRIVER_LICENSE": "US_DL",
    "UK_NHS": "UK_NHS", "UK_NIN": "UK_NIN", "UK_PASSPORT": "UK_PASSPORT",
    "UK_CRN": "UK_CRN", "UK_DRIVING_LICENCE": "UK_DL",
    "EU_VAT": "EU_VAT", "EU_PASSPORT": "EU_PASSPORT",
    "DE_TAX_ID": "DE_TAX", "DE_SOCIAL_SECURITY": "DE_SSN",
    "FR_NIR": "FR_NIR", "FR_CNI": "FR_CNI",
    "IT_FISCAL_CODE": "IT_CF", "IT_VAT": "IT_VAT",
    "ES_DNI": "ES_DNI", "ES_NIE": "ES_NIE",
    "CY_TIC": "CY_TIC", "CY_ID_CARD": "CY_ID",
    "MEDICAL_LICENSE": "MED_LIC",
}


# ============================================================
# Mapping persistence + TTL
# ============================================================
_in_memory_mappings = {}

def save_mapping(session_id, mapping, metadata=None):
    data = {"session_id": session_id, "mapping": mapping, "metadata": metadata or {}, "timestamp": time.time()}
    # Always keep in memory FIRST — this never fails
    _in_memory_mappings[session_id] = data
    # Try to persist to disk (optional — in-memory is the primary store)
    disk_path = None
    try:
        MAPPING_DIR.mkdir(parents=True, exist_ok=True)
        path = MAPPING_DIR / f"{session_id}.json"
        tmp = path.with_suffix(".tmp")
        tmp.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
        tmp.replace(path)
        disk_path = str(path)
    except Exception as e:
        log.warning(f"save_mapping disk write failed (in-memory OK): {e}")
    return disk_path or f"memory://{session_id}"


def load_mapping(session_id):
    # Try disk first
    try:
        path = MAPPING_DIR / f"{session_id}.json"
        if path.exists():
            return json.loads(path.read_text(encoding="utf-8")).get("mapping", {})
    except Exception as e:
        log.warning(f"load_mapping disk read failed: {e}")
    # Fallback to in-memory
    if session_id in _in_memory_mappings:
        return _in_memory_mappings[session_id].get("mapping", {})
    return {}


def cleanup_old_mappings():
    """Delete mappings older than TTL."""
    cutoff = time.time() - (MAPPING_TTL_DAYS * 86400)
    removed = 0
    for f in MAPPING_DIR.glob("*.json"):
        try:
            if f.stat().st_mtime < cutoff:
                f.unlink()
                removed += 1
        except Exception:
            pass
    if removed:
        log.info(f"Cleaned up {removed} expired mappings (>{MAPPING_TTL_DAYS} days)")


# ============================================================
# PIIEngine
# ============================================================
class PIIEngine:
    _instance = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._initialized = False
        return cls._instance

    GLINER_MODEL_NAME = os.environ.get("PII_GLINER_MODEL", "urchade/gliner_large-v2.1")

    def _ensure_ready(self, _from_bootstrap=False):
        """Initialize the PII engine. Called from bootstrap thread after packages/models are ready,
        or lazily on first tool call.

        Uses GLiNER (DeBERTa-v3 zero-shot NER) for high-quality entity recognition.
        SpaCy handles tokenization, GLiNER handles NER via Presidio's GLiNERRecognizer.
        Falls back to SpaCy-only if GLiNER is unavailable.
        """
        if self._initialized:
            return

        # If called from a tool (not from bootstrap), wait for bootstrap to finish first.
        if not _from_bootstrap:
            global _bootstrap_done
            if not _bootstrap_done:
                log.info("Waiting for background bootstrap to complete...")
                for _ in range(600):
                    if _bootstrap_done:
                        break
                    time.sleep(1)
                if not _bootstrap_done:
                    raise RuntimeError(
                        "Bootstrap timed out after 10 minutes. "
                        "Check internet connection and try restarting."
                    )
            if _bootstrap_error:
                log.warning(f"Bootstrap had errors: {_bootstrap_error}")

        log.info("Initializing PII Engine v5.4.0 (lazy init on first use)...")

        # --- SpaCy NLP engine (tokenization only) ---
        from presidio_analyzer.nlp_engine import NlpEngineProvider
        try:
            nlp_engine = NlpEngineProvider(nlp_configuration={
                "nlp_engine_name": "spacy",
                "models": [{"lang_code": "en", "model_name": "en_core_web_sm"}],
            }).create_engine()
        except Exception as e:
            raise RuntimeError(f"SpaCy engine failed: {e}. Ensure spacy and en_core_web_sm are installed.") from e

        # --- Build registry with built-in recognizers ---
        from presidio_analyzer import AnalyzerEngine, RecognizerRegistry
        registry = RecognizerRegistry()
        registry.load_predefined_recognizers()
        log.info(f"Loaded {len(registry.recognizers)} predefined recognizers")

        # --- Try GLiNER (zero-shot NER), fallback to SpaCy-only ---
        backend_used = "spacy (en_core_web_sm) [FALLBACK]"
        try:
            from presidio_analyzer.predefined_recognizers import GLiNERRecognizer
            gliner_recognizer = GLiNERRecognizer(
                model_name=self.GLINER_MODEL_NAME,
                entity_mapping={
                    "person": "PERSON",
                    "company": "ORGANIZATION",
                    "organization": "ORGANIZATION",
                    "location": "LOCATION",
                    "nationality": "NRP",
                },
                flat_ner=False,
                multi_label=True,
                map_location="cpu",
            )
            registry.add_recognizer(gliner_recognizer)
            backend_used = f"gliner ({self.GLINER_MODEL_NAME})"
            log.info(f"GLiNER recognizer loaded: {backend_used}")
        except Exception as e:
            log.warning(f"GLiNER failed: {e}")
            log.warning("Using SpaCy-only NER (reduced quality)")

        self.analyzer = AnalyzerEngine(nlp_engine=nlp_engine, registry=registry)
        self._backend = backend_used

        # EU recognizers
        try:
            import importlib.util
            eu_path = Path(__file__).parent / "eu_recognizers.py"
            if eu_path.exists():
                spec = importlib.util.spec_from_file_location("eu_recognizers", eu_path)
                mod = importlib.util.module_from_spec(spec)
                spec.loader.exec_module(mod)
                count = mod.register_eu_recognizers(self.analyzer)
                log.info(f"EU recognizers: {count} registered")
        except Exception as e:
            log.warning(f"EU recognizers failed: {e}")

        self._initialized = True

        cleanup_old_mappings()
        log.info(f"PII Engine ready ({backend_used})")

    def __init__(self):
        # Lightweight — just mark as not initialized.
        # Heavy work is deferred to _ensure_ready() on first tool call.
        pass

    # --- Core detect ---
    def _deduplicate(self, results):
        if not results:
            return []
        s = sorted(results, key=lambda r: (r.start, -r.score))
        d = [s[0]]
        for r in s[1:]:
            if r.start >= d[-1].end:
                d.append(r)
            elif r.score > d[-1].score:
                d[-1] = r
        return d

    # Entity types that are proper-noun based (names, orgs, locations)
    _NAMED_ENTITY_TYPES = {"PERSON", "ORGANIZATION", "LOCATION", "NRP"}

    @staticmethod
    def _snap_word_boundaries(text, entities):
        """Pass 1: Snap entity boundaries to word edges.

        If NER cut a word in the middle, extend to complete that word.
        Never adds new words — only finishes the word already partially captured.
        Never snaps across newlines (paragraph boundaries).
        Also trims trailing/leading punctuation.
        Drops entities that are too short (<=2 chars) after cleanup.
        """
        tlen = len(text)
        _split_buf = []  # collect newline-split entities here (not inside loop)
        for e in entities:
            start, end = e["start"], e["end"]

            # Snap RIGHT: if boundary is mid-word, complete the word
            # Stop at newlines — don't merge across paragraphs
            if end < tlen and end > 0 and text[end].isalnum() and text[end - 1].isalnum():
                while end < tlen and text[end].isalnum() and text[end] != '\n':
                    end += 1

            # Snap LEFT: if boundary is mid-word, complete the word
            if start > 0 and start < end and text[start].isalnum() and text[start - 1].isalnum():
                while start > 0 and text[start - 1].isalnum() and text[start - 1] != '\n':
                    start -= 1

            # Trim trailing/leading punctuation, whitespace, newlines
            while end > start and text[end - 1] in '.,;:)]\'" \t\n\r':
                end -= 1
            while start < end and text[start] in '([\'" \t\n\r#/':
                start += 1

            entity_text = text[start:end].strip()
            if len(entity_text) <= 2:
                # Too short to be meaningful PII (e.g. "S", "St")
                e["_drop"] = True
                log.info(f"Boundary drop (too short): '{entity_text}' "
                         f"(type={e.get('type', '?')})")
            elif '\n' in entity_text:
                # Entity spans multiple lines — split into separate entities per line
                # NOTE: collect into _split_buf to avoid mutating list during iteration
                e["_drop"] = True
                lines = entity_text.split('\n')
                search_from = start
                for line in lines:
                    stripped = line.strip()
                    if len(stripped) > 2:
                        line_start = text.find(stripped, search_from)
                        if line_start == -1:
                            continue  # skip if not found (safety)
                        _split_buf.append({
                            "start": line_start,
                            "end": line_start + len(stripped),
                            "text": stripped,
                            "type": e.get("type", ""),
                            "score": e.get("score", 0),
                        })
                        search_from = line_start + len(stripped)
                    else:
                        # Advance search position past this line
                        pos = text.find(line, search_from)
                        if pos != -1:
                            search_from = pos + len(line)
                log.info(f"Boundary split (across newline): '{entity_text[:40]}' → "
                         f"{[l.strip() for l in lines if len(l.strip()) > 2]}")
            elif start < end:
                e["start"] = start
                e["end"] = end
                e["text"] = entity_text
            else:
                e["_drop"] = True

        # Append split entities collected during the loop (avoids list mutation during iteration)
        entities.extend(_split_buf)
        return [e for e in entities if not e.get("_drop")]

    # Common legal/contract terms that NER frequently misclassifies as PII.
    # Case-insensitive match against normalized entity text.
    _LEGAL_STOPLIST = {
        # Contract parties / roles
        "contractor", "client", "customer", "vendor", "supplier",
        "licensor", "licensee", "employer", "employee", "consultant",
        "subcontractor", "agent", "principal", "assignee", "assignor",
        "guarantor", "beneficiary", "trustee", "grantor", "grantee",
        "lessee", "lessor", "tenant", "landlord", "borrower", "lender",
        "buyer", "seller", "partner", "shareholder", "director",
        "officer", "secretary", "treasurer", "representative",
        # Document terms
        "order", "agreement", "contract", "amendment", "addendum",
        "exhibit", "schedule", "appendix", "annex", "section",
        "article", "clause", "paragraph", "party", "parties",
        "purchase order", "statement of work", "scope of work",
        # Generic terms NER often catches
        "company", "corporation", "entity", "firm", "business",
        "affiliate", "subsidiary", "parent", "division", "branch",
        "effective date", "termination date", "commencement date",
        # Software / product / brand names (not PII)
        "adobe", "adobe premiere", "adobe premiere pro", "adobe after effects",
        "final cut", "final cut pro", "davinci resolve",
        "photoshop", "illustrator", "figma", "canva",
        "microsoft", "google", "apple", "amazon", "meta",
        # Common with Cyrillic homoglyphs (С = Cyrillic Es looks like Latin C)
        "\u0441lient", "сlient",  # Cyrillic С + lient
    }

    @staticmethod
    def _filter_false_positives(entities):
        """Pass 2: Drop false positives using cross-entity context and stop-list.

        Rules:
        1. Stop-list: known legal/contract terms → always drop.
        2. Single lowercase word + named entity type → drop (not a proper noun).
        3. If same text appears in another entity with higher score → confirms both.
        """
        # Collect all confirmed high-score entity texts for cross-reference
        confirmed_texts = set()
        for e in entities:
            if e.get("score", 0) >= 0.6:
                confirmed_texts.add(e["text"].lower())
                for word in e["text"].split():
                    confirmed_texts.add(word.lower())

        cleaned = []
        for e in entities:
            txt = e["text"]
            etype = e.get("type", "")
            words = txt.split()
            norm_txt = txt.lower().strip()

            # Rule 0: Stop-list — known non-PII terms
            # Also check with Cyrillic homoglyph normalization (С→C, А→A, etc.)
            _CYRILLIC_TO_LATIN = str.maketrans('СсАаЕеОоРрХхВвМмТтНн', 'CcAaEeOoPpXxBbMmTtHh')
            norm_latin = norm_txt.translate(_CYRILLIC_TO_LATIN)
            if norm_txt in PIIEngine._LEGAL_STOPLIST or norm_latin in PIIEngine._LEGAL_STOPLIST:
                log.info(f"FP drop (stop-list): '{txt}' (type={etype})")
                continue

            # Rule 1: Single lowercase word + named entity type → likely false positive
            if len(words) == 1 and etype in PIIEngine._NAMED_ENTITY_TYPES:
                if txt[0].islower():
                    if txt.lower() not in confirmed_texts:
                        log.info(f"FP drop (single lowercase word): '{txt}' "
                                 f"(type={etype}, score={e.get('score', '?')})")
                        continue

            cleaned.append(e)
        return cleaned

    @classmethod
    def _clean_boundaries(cls, text, entities):
        """Two-pass boundary cleanup: snap words, then filter false positives."""
        entities = cls._snap_word_boundaries(text, entities)
        entities = cls._filter_false_positives(entities)
        return entities

    def _analyze_chunked(self, text, language="en", chunk_size=800, overlap=100):
        """Run analyzer on text in chunks to avoid GLiNER timeout on long texts.
        Chunks overlap to avoid splitting entities at boundaries."""
        if len(text) <= chunk_size:
            return self.analyzer.analyze(text=text, entities=SUPPORTED_ENTITIES, language=language)

        all_results = []
        start = 0
        while start < len(text):
            end = min(start + chunk_size, len(text))
            # Try to break at whitespace
            if end < len(text):
                ws = text.rfind(' ', start + chunk_size - overlap, end)
                if ws > start:
                    end = ws + 1
            chunk = text[start:end]
            chunk_results = self.analyzer.analyze(text=chunk, entities=SUPPORTED_ENTITIES, language=language)
            # Adjust offsets to full text positions
            for r in chunk_results:
                r.start += start
                r.end += start
                all_results.append(r)
            start = end - overlap if end < len(text) else len(text)

        # Deduplicate overlapping detections (same span)
        seen = set()
        unique = []
        for r in sorted(all_results, key=lambda x: (x.start, -x.score)):
            key = (r.start, r.end, r.entity_type)
            if key not in seen:
                seen.add(key)
                unique.append(r)
        return unique

    def detect(self, text, language="en"):
        """Detect PII using NER. All entities above MIN_SCORE are confirmed."""
        self._ensure_ready()
        results = self._analyze_chunked(text, language)
        results = self._deduplicate(results)

        entities = []

        for r in results:
            et = text[r.start:r.end]
            etype = r.entity_type

            if r.score < MIN_SCORE:
                continue

            entry = {
                "text": et, "type": etype, "start": r.start, "end": r.end,
                "score": round(r.score, 3), "verified": True, "reason": "NER",
            }
            entities.append(entry)

        entities = self._clean_boundaries(text, entities)
        log.info(f"Detect: {len(entities)} entities confirmed (NER-only)")
        return entities

    # --- Fuzzy entity deduplication helpers ---
    @staticmethod
    def _normalize(text):
        """Normalize entity text for dedup: lowercase, strip punctuation, collapse spaces."""
        return re.sub(r'\s+', ' ', text.lower().strip().rstrip('.,;:'))

    def _get_or_create_placeholder(self, etype, text, type_counters, seen_exact, seen_family, mapping, prefix=""):
        """Get existing placeholder for this exact entity text, or create new one.

        Exact-match dedup: "Acme" always maps to the same placeholder.
        Family grouping: "Acme" → <ORG_1>, "Acme Corp." → <ORG_1a>,
        "Acme Corporation" → <ORG_1b>. Each has its own placeholder
        with its own reverse mapping, so deanonymization restores exact text.

        Args:
            etype: entity type string
            text: raw entity text as found in document
            type_counters: dict tracking next family index per type
            seen_exact: dict of (type, exact_normalized_text) → placeholder
            seen_family: dict of (type, short_norm) → (family_number, variant_counter)
            mapping: dict of placeholder → raw text
            prefix: optional prefix for multi-file workflows
        Returns:
            placeholder string
        """
        norm = self._normalize(text)
        exact_key = (etype, norm)

        # 1. Exact match — same text seen before, reuse placeholder
        if exact_key in seen_exact:
            return seen_exact[exact_key]

        tag = TAG_NAMES.get(etype, etype)

        # 2. Check if this belongs to an existing family (substring match)
        family_key = None
        if len(norm) >= 4:
            for (ft, fn), (fnum, _) in seen_family.items():
                if ft != etype:
                    continue
                if len(fn) >= 4 and (norm in fn or fn in norm):
                    family_key = (ft, fn)
                    break

        if family_key:
            # Add as variant to existing family
            fnum, vcounter = seen_family[family_key]
            vcounter += 1
            seen_family[family_key] = (fnum, vcounter)
            # variant suffix: a, b, c, ...
            suffix = chr(ord('a') + vcounter - 1) if vcounter <= 26 else str(vcounter)
            if prefix:
                placeholder = f"<{prefix}_{tag}_{fnum}{suffix}>"
            else:
                placeholder = f"<{tag}_{fnum}{suffix}>"
        else:
            # New family
            type_counters[etype] += 1
            fnum = type_counters[etype]
            if prefix:
                placeholder = f"<{prefix}_{tag}_{fnum}>"
            else:
                placeholder = f"<{tag}_{fnum}>"
            # Register this as the family root (shortest/first form)
            seen_family[(etype, norm)] = (fnum, 0)

        seen_exact[exact_key] = placeholder
        mapping[placeholder] = text  # exact text for this specific placeholder
        log.info(f"Placeholder: '{text}' → {placeholder}")
        return placeholder

    # --- Assign indexed placeholders (shared logic) ---
    def _assign_placeholders(self, confirmed_entities, prefix=""):
        """Assign indexed placeholders preserving exact entity forms. Returns mapping dict."""
        type_counters = defaultdict(int)
        seen_exact = {}   # (type, exact_normalized_text) → placeholder
        seen_family = {}  # (type, normalized_text) → (family_number, variant_counter)
        mapping = {}      # placeholder → exact raw text

        for e in sorted(confirmed_entities, key=lambda x: x["start"]):
            e["placeholder"] = self._get_or_create_placeholder(
                e["type"], e["text"], type_counters, seen_exact, seen_family, mapping, prefix
            )

        return mapping

    # --- Text anonymization ---
    def anonymize_text(self, text, language="en", prefix=""):
        t0 = time.time()
        entities = self.detect(text, language)
        confirmed = [e for e in entities if e.get("verified")]

        mapping = self._assign_placeholders(confirmed, prefix)

        anonymized = text
        for e in sorted(confirmed, key=lambda x: x["start"], reverse=True):
            anonymized = anonymized[:e["start"]] + e["placeholder"] + anonymized[e["end"]:]

        session_id = uuid.uuid4().hex[:12]
        save_mapping(session_id, mapping, {"confirmed": len(confirmed)})

        by_type = defaultdict(int)
        for e in confirmed:
            by_type[e["type"]] += 1

        # Build safe entity list — strip real text, keep only placeholders and metadata
        safe_entities = []
        for e in confirmed:
            safe_entities.append({
                "placeholder": e.get("placeholder", ""),
                "type": e["type"],
                "score": e["score"],
                "verified": e["verified"],
                "reason": e.get("reason", ""),
            })

        return {
            "anonymized_text": anonymized, "session_id": session_id,
            "total_entities": len(entities), "entities_confirmed": len(confirmed),
            "unique_entities": len(mapping),
            "by_type": dict(by_type), "entities": safe_entities,
            "processing_time_ms": round((time.time() - t0) * 1000, 1),
        }

    # --- DOCX anonymization (reuses detect) ---
    def anonymize_docx(self, docx_path, language="en", prefix=""):
        from docx import Document
        t0 = time.time()
        doc = Document(str(docx_path))

        # Shared state across paragraphs (exact dedup + family grouping)
        type_counters = defaultdict(int)
        seen_exact = {}   # (type, exact_normalized_text) → placeholder
        seen_family = {}  # (type, normalized_text) → (family_number, variant_counter)
        mapping = {}      # placeholder → exact raw text
        total = 0
        by_type = defaultdict(int)

        for para in self._iter_docx_paragraphs(doc):
            full_text, runs_info = self._get_runs(para)
            if not full_text.strip():
                continue

            entities = self.detect(full_text, language)
            confirmed = [e for e in entities if e.get("verified")]

            for e in sorted(confirmed, key=lambda x: x["start"], reverse=True):
                ph = self._get_or_create_placeholder(
                    e["type"], e["text"], type_counters, seen_exact, seen_family, mapping, prefix
                )
                self._replace_in_runs(runs_info, e["start"], e["end"], ph)
                total += 1
                by_type[e["type"]] += 1

        docx_path = Path(docx_path)
        out_path = docx_path.parent / f"{docx_path.stem}_anonymized.docx"
        doc.save(str(out_path))

        session_id = uuid.uuid4().hex[:12]
        save_mapping(session_id, mapping, {"source": str(docx_path)})

        return {
            "output_path": str(out_path), "session_id": session_id,
            "total_entities": total,
            "unique_entities": len(mapping), "by_type": dict(by_type),
            "processing_time_ms": round((time.time() - t0) * 1000, 1),
        }

    # --- Deanonymization ---
    @staticmethod
    def deanonymize_text(text, mapping):
        for ph in sorted(mapping.keys(), key=len, reverse=True):
            text = text.replace(ph, mapping[ph])
        return text

    def deanonymize_docx(self, docx_path, mapping):
        from docx import Document
        doc = Document(str(docx_path))
        sorted_ph = sorted(mapping.keys(), key=len, reverse=True)
        for para in self._iter_docx_paragraphs(doc):
            for run in para.runs:
                t = run.text
                for ph in sorted_ph:
                    if ph in t:
                        t = t.replace(ph, mapping[ph])
                run.text = t
        out = Path(docx_path).parent / f"{Path(docx_path).stem}_restored.docx"
        doc.save(str(out))
        return str(out)

    # --- DOCX helpers ---
    @staticmethod
    def _iter_docx_paragraphs(doc):
        for p in doc.paragraphs:
            yield p
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
        for sec in doc.sections:
            for hf in [sec.header, sec.footer]:
                if hf:
                    for p in hf.paragraphs:
                        yield p

    @staticmethod
    def _get_runs(para):
        runs_info = []
        offset = 0
        for run in para.runs:
            runs_info.append({"run": run, "text": run.text, "start": offset, "end": offset + len(run.text)})
            offset += len(run.text)
        return "".join(r["text"] for r in runs_info), runs_info

    @staticmethod
    def _replace_in_runs(runs_info, start, end, replacement):
        """Replace text at [start, end) with replacement across runs.

        Uses a two-phase approach: modify run texts first, then recalculate
        all offsets from scratch. This avoids cascading offset errors when
        an entity spans multiple runs.
        """
        # Phase 1: identify affected runs and modify their text
        affected = []
        for i, ri in enumerate(runs_info):
            if ri["end"] <= start or ri["start"] >= end:
                continue
            affected.append(i)

        if not affected:
            return

        first_idx = affected[0]
        last_idx = affected[-1]

        for idx in affected:
            ri = runs_info[idx]
            old_text = ri["text"]
            local_start = max(0, start - ri["start"])
            local_end = min(len(old_text), end - ri["start"])

            if idx == first_idx and idx == last_idx:
                # Entity fully within one run
                new_text = old_text[:local_start] + replacement + old_text[local_end:]
            elif idx == first_idx:
                # Entity starts here, continues into next run(s)
                new_text = old_text[:local_start] + replacement
            elif idx == last_idx:
                # Entity ends here, started in previous run(s)
                new_text = old_text[local_end:]
            else:
                # Middle run — entirely within entity range
                new_text = ""

            ri["run"].text = new_text
            ri["text"] = new_text

        # Phase 2: recalculate all offsets from scratch
        offset = 0
        for ri in runs_info:
            ri["start"] = offset
            ri["end"] = offset + len(ri["text"])
            offset += len(ri["text"])


# ============================================================
# MCP Tools
# ============================================================
mcp = FastMCP("PII Shield", host="127.0.0.1", port=int(os.environ.get("PII_PORT", "8765")))
engine = PIIEngine()


def _latest_session_id():
    """Find the most recent session by mapping file mtime, with in-memory fallback."""
    try:
        sessions = sorted(MAPPING_DIR.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
        if sessions:
            return json.loads(sessions[0].read_text(encoding="utf-8")).get("session_id", "")
    except Exception:
        pass
    # Fallback to in-memory
    if _in_memory_mappings:
        latest = max(_in_memory_mappings.values(), key=lambda d: d.get("timestamp", 0))
        return latest.get("session_id", "")
    return ""


def _check_ready():
    """Check if bootstrap is done. Returns None if ready, or a JSON status string if still loading."""
    if _bootstrap_done and not _bootstrap_error:
        return None  # Ready
    if _bootstrap_error and _bootstrap_phase == "error":
        return json.dumps({
            "status": "error",
            "message": f"PII Shield failed to initialize: {_bootstrap_error}",
            "hint": "Check internet connection, ensure Python 3.10+ with pip, and restart Cowork.",
        }, indent=2)
    # Still loading — include progress info
    elapsed = round(time.time() - _bootstrap_start_time, 1) if _bootstrap_start_time else 0
    progress_map = {"starting": 5, "packages": 40, "models": 70, "engine": 90}
    progress_pct = progress_map.get(_bootstrap_phase, 0)
    return json.dumps({
        "status": "loading",
        "phase": _bootstrap_phase,
        "message": _bootstrap_detail or "PII Shield is starting up...",
        "progress_pct": progress_pct,
        "elapsed_seconds": elapsed,
        "hint": "First-time setup installs dependencies (~5-10 min). Please wait and try again.",
    }, indent=2)


@mcp.tool()
def anonymize_text(text: str, language: str = "en", prefix: str = "") -> str:
    """Anonymize PII in text. Returns indexed placeholders + session_id for deanonymization.
    Use prefix (e.g. "D1") for multi-file workflows to avoid placeholder collisions."""
    loading = _check_ready()
    if loading:
        return loading
    r = engine.anonymize_text(text, language, prefix=prefix)
    return json.dumps(r, indent=2, ensure_ascii=False)


@mcp.tool()
def anonymize_file(file_path: str, language: str = "en", prefix: str = "") -> str:
    """Anonymize PII in a file. Auto-detects format: .docx (preserves formatting), .txt/.md/.csv (plain text).
    Use prefix (e.g. "D1") for multi-file workflows to avoid placeholder collisions."""
    loading = _check_ready()
    if loading:
        return loading
    p = Path(file_path).expanduser().resolve()
    if not p.exists():
        return json.dumps({"error": f"File not found: {p}"})

    if p.suffix.lower() == ".docx":
        r = engine.anonymize_docx(p, language, prefix=prefix)
        return json.dumps(r, indent=2, ensure_ascii=False)
    elif p.suffix.lower() in (".txt", ".md", ".csv", ".log", ".text"):
        text = p.read_text(encoding="utf-8")
        r = engine.anonymize_text(text, language, prefix=prefix)
        out = p.parent / f"{p.stem}_anonymized{p.suffix}"
        out.write_text(r["anonymized_text"], encoding="utf-8")
        r["output_path"] = str(out)
        return json.dumps(r, indent=2, ensure_ascii=False)
    else:
        return json.dumps({"error": f"Unsupported format: {p.suffix}. Supported: .docx .txt .md .csv"})


@mcp.tool()
def anonymize_docx(file_path: str, language: str = "en", prefix: str = "") -> str:
    """Anonymize PII in .docx preserving all formatting. Use for round-trip document editing.
    Use prefix (e.g. "D1") for multi-file workflows to avoid placeholder collisions."""
    loading = _check_ready()
    if loading:
        return loading
    p = Path(file_path).expanduser().resolve()
    if not p.exists():
        return json.dumps({"error": f"Not found: {p}"})
    r = engine.anonymize_docx(p, language, prefix=prefix)
    return json.dumps(r, indent=2, ensure_ascii=False)


@mcp.tool()
def deanonymize_text(text: str, session_id: str = "", output_path: str = "") -> str:
    """Restore real PII values in text. Writes result to .docx file — never returns PII to Claude.
    Returns only the file path. output_path should end with .docx (default) or .txt."""
    sid = session_id.strip() or _latest_session_id()
    if not sid:
        return json.dumps({"error": "No session. Run anonymize first."})
    mapping = load_mapping(sid)
    if not mapping:
        return json.dumps({"error": f"Mapping not found: {sid}"})
    restored = engine.deanonymize_text(text, mapping)
    # Determine output path
    if output_path:
        out_p = Path(output_path).expanduser().resolve()
    else:
        out_p = MAPPING_DIR / f"restored_{sid}.docx"
    out_p.parent.mkdir(parents=True, exist_ok=True)
    # Write to file — PII never goes back to Claude
    if out_p.suffix.lower() == ".docx":
        _write_docx(restored, out_p)
    else:
        out_p.write_text(restored, encoding="utf-8")
    return json.dumps({
        "restored_path": str(out_p),
        "session_id": sid,
        "entities_restored": len(mapping),
        "note": "PII-safe: restored text written to file, not returned to LLM.",
    }, indent=2, ensure_ascii=False)


def _write_docx(text: str, path: Path):
    """Write text to a formatted .docx file."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
    lines = text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        # Detect headings by common patterns
        if stripped.isupper() and len(stripped) < 100:
            p = doc.add_paragraph(stripped)
            p.style = doc.styles["Heading 2"]
        elif stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            p = doc.add_paragraph(stripped.lstrip("# "))
            p.style = doc.styles[f"Heading {level}"]
        else:
            doc.add_paragraph(stripped)
    doc.save(str(path))


@mcp.tool()
def deanonymize_docx(file_path: str, session_id: str = "") -> str:
    """Restore real PII in .docx preserving formatting."""
    sid = session_id.strip() or _latest_session_id()
    if not sid:
        return json.dumps({"error": "No session. Run anonymize first."})
    mapping = load_mapping(sid)
    if not mapping:
        return json.dumps({"error": f"Mapping not found: {sid}"})
    p = Path(file_path).expanduser().resolve()
    if not p.exists():
        return json.dumps({"error": f"Not found: {p}"})
    out = engine.deanonymize_docx(p, mapping)
    return json.dumps({"restored_path": out, "session_id": sid}, indent=2, ensure_ascii=False)


@mcp.tool()
def get_mapping(session_id: str = "") -> str:
    """Retrieve mapping metadata (placeholder keys and entity types only — no real PII values).
    Full mapping stays on disk, never returned to LLM."""
    sid = session_id.strip() or _latest_session_id()
    if not sid:
        return json.dumps({"error": "No session available."})
    mapping = load_mapping(sid)
    if not mapping:
        return json.dumps({"error": f"Not found: {sid}"})
    # Return only placeholder keys and types — no real PII values
    safe_summary = {}
    for placeholder in mapping:
        # Extract type from placeholder like <PERSON_1> → PERSON
        etype = placeholder.strip("<>").rsplit("_", 1)[0] if "_" in placeholder else placeholder.strip("<>")
        safe_summary[placeholder] = etype
    return json.dumps({
        "session_id": sid,
        "total_entities": len(mapping),
        "placeholders": safe_summary,
        "note": "PII-safe: real values not returned. Use deanonymize_text/docx to restore to file.",
    }, indent=2, ensure_ascii=False)


@mcp.tool()
def scan_text(text: str, language: str = "en") -> str:
    """Detect PII without anonymizing. Preview mode. Returns entity types and positions, not real text values."""
    loading = _check_ready()
    if loading:
        return loading
    entities = engine.detect(text, language)
    # Strip real text from entities — return only type, position, and verification status
    safe_entities = []
    for e in entities:
        safe_entities.append({
            "type": e["type"],
            "start": e["start"],
            "end": e["end"],
            "score": e["score"],
            "verified": e.get("verified"),
            "reason": e.get("reason", ""),
        })
    return json.dumps({
        "entities_found": len(entities),
        "confirmed": len([e for e in entities if e.get("verified")]),
        "rejected": len([e for e in entities if not e.get("verified")]),
        "entities": safe_entities,
    }, indent=2, ensure_ascii=False)


@mcp.tool()
def list_entities() -> str:
    """Show status, supported types, and recent sessions."""
    # Always show recent sessions (no engine needed)
    sessions = sorted(MAPPING_DIR.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)[:5]
    recent = []
    for s in sessions:
        try:
            d = json.loads(s.read_text())
            recent.append({"session_id": d["session_id"], "entities": len(d["mapping"])})
        except Exception:
            pass

    # If still bootstrapping, show status without crashing
    if not _bootstrap_done:
        return json.dumps({
            "status": "loading",
            "phase": _bootstrap_phase,
            "message": _bootstrap_detail or "PII Shield is starting up...",
            "hint": "First-time setup installs dependencies (~5-10 min). Please wait and try again.",
            "recent_sessions": recent,
        }, indent=2, ensure_ascii=False)

    if _bootstrap_error and _bootstrap_phase == "error":
        return json.dumps({
            "status": "error",
            "message": f"Bootstrap failed: {_bootstrap_error}",
            "recent_sessions": recent,
        }, indent=2, ensure_ascii=False)

    # Engine ready — full diagnostics
    eng = PIIEngine()
    eng._ensure_ready()
    recognizer_names = [type(r).__name__ for r in eng.analyzer.registry.recognizers]
    engine_class = type(eng.analyzer.nlp_engine).__name__

    backend = eng._backend or ""
    if "gliner" in backend:
        quality = "full (GLiNER zero-shot)"
    elif "FALLBACK" not in backend:
        quality = "full"
    else:
        quality = "reduced (SpaCy fallback)"
    return json.dumps({
        "status": "ready",
        "mode": "NER-only",
        "backend": eng._backend,
        "quality": quality,
        "nlp_engine_class": engine_class,
        "recognizers": recognizer_names,
        "min_score": MIN_SCORE,
        "mapping_ttl_days": MAPPING_TTL_DAYS,
        "recent_sessions": recent,
    }, indent=2, ensure_ascii=False)


def _ensure_ssl_cert(cert_dir: Path):
    """Generate self-signed cert if not exists."""
    cert_file = cert_dir / "cert.pem"
    key_file = cert_dir / "key.pem"
    if cert_file.exists() and key_file.exists():
        return str(cert_file), str(key_file)
    cert_dir.mkdir(parents=True, exist_ok=True)
    log.info("Generating self-signed SSL certificate...")
    import subprocess
    subprocess.run([
        "python", "-c",
        f"""
from cryptography import x509
from cryptography.x509.oid import NameOID
from cryptography.hazmat.primitives import hashes, serialization
from cryptography.hazmat.primitives.asymmetric import rsa
import datetime, ipaddress
key = rsa.generate_private_key(public_exponent=65537, key_size=2048)
name = x509.Name([x509.NameAttribute(NameOID.COMMON_NAME, "localhost")])
cert = (x509.CertificateBuilder()
    .subject_name(name).issuer_name(name)
    .public_key(key.public_key())
    .serial_number(x509.random_serial_number())
    .not_valid_before(datetime.datetime.now(datetime.timezone.utc))
    .not_valid_after(datetime.datetime.now(datetime.timezone.utc) + datetime.timedelta(days=3650))
    .add_extension(x509.SubjectAlternativeName([
        x509.DNSName("localhost"),
        x509.IPAddress(ipaddress.IPv4Address("127.0.0.1")),
    ]), critical=False)
    .sign(key, hashes.SHA256()))
open(r"{cert_file}", "wb").write(cert.public_bytes(serialization.Encoding.PEM))
open(r"{key_file}", "wb").write(key.private_bytes(
    serialization.Encoding.PEM, serialization.PrivateFormat.TraditionalOpenSSL, serialization.NoEncryption()))
print("OK")
"""
    ], check=True)
    log.info(f"SSL cert created: {cert_file}")
    return str(cert_file), str(key_file)


if __name__ == "__main__":
    transport = "sse" if "--sse" in sys.argv else "stdio"
    port = int(os.environ.get("PII_PORT", "8765"))

    log.info(f"Starting PII Shield MCP Server v5.4.0 ({transport})...")

    if transport == "sse":
            import ssl
            import uvicorn

            cert_dir = Path.home() / ".pii-shield" / "ssl"
            cert_file, key_file = _ensure_ssl_cert(cert_dir)

            app = mcp.sse_app()
            uvicorn.run(
                app,
                host="127.0.0.1",
                port=port,
                ssl_certfile=cert_file,
                ssl_keyfile=key_file,
                log_level="info",
            )
    else:
        mcp.run(transport="stdio")
